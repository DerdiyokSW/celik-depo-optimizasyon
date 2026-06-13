"""GEÇİCİ — BITIRME-RAPORU.md'yi Times New Roman bir .docx'e çevirir (şablona uygun).

Başlıklar gerçek Word Heading stillerine atanır (otomatik İçindekiler için). İÇİNDEKİLER
başlığının altına canlı TOC alanı gömülür. ![altyazı](png) satırları şekil + altyazı olarak
eklenir. Kapak başlık bloğu ortalanır. Çıktı: BITIRME-RAPORU.docx
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(".").resolve()
SRC = ROOT / "BITIRME-RAPORU.md"
OUT = ROOT / "BITIRME-RAPORU.docx"
FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
FRONT_BACK = {"ÖNSÖZ", "İÇİNDEKİLER", "SİMGELER VE KISALTMALAR LİSTESİ", "ŞEKİLLER LİSTESİ",
              "TABLOLAR LİSTESİ", "ÖZET", "KAYNAKLAR", "EKLER", "ÖZGEÇMİŞ"}
INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+?`|\*[^*\n]+?\*)")
IMG = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
# Heading stillerini TNR + siyah yap (İçindekiler de TNR olsun).
for lvl, sz in ((1, 14), (2, 13), (3, 12)):
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = FONT
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = BLACK


def style_run(r, size=12, bold=False, italic=False, mono=False):
    r.font.name = FONT if not mono else "Consolas"
    r.font.size = Pt(size)
    if bold:
        r.bold = True
    if italic:
        r.italic = True


def add_inline(p, text, size=12):
    text = text.replace("&nbsp;", " ")
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            style_run(p.add_run(part[2:-2]), size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            style_run(p.add_run(part[1:-1]), size, mono=True)
        elif part.startswith("*") and part.endswith("*"):
            style_run(p.add_run(part[1:-1]), size, italic=True)
        else:
            style_run(p.add_run(part), size)


def add_toc():
    """Word'ün açılışta dolduracağı canlı İçindekiler alanı."""
    p = doc.add_paragraph()
    r = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = r'TOC \o "1-3" \h \z \u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "İçindekiler — güncellemek için seçip F9'a basın."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for e in (f1, it, f2, t, f3):
        r._r.append(e)


def add_picture(path, caption):
    full = ROOT / path
    if not full.exists():
        p = doc.add_paragraph(); style_run(p.add_run(f"[Şekil bulunamadı: {path}]"), italic=True)
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(full), width=Inches(6.0))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    style_run(cap.add_run(caption), size=10, italic=True)


def add_table(rows):
    body = [r for r in rows if not re.match(r"^[:\-\| ]+$", "|".join(r))]
    if not body:
        return
    ncol = max(len(r) for r in body)
    tbl = doc.add_table(rows=0, cols=ncol)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(body):
        cells = tbl.add_row().cells
        for ci in range(ncol):
            par = cells[ci].paragraphs[0]
            add_inline(par, row[ci] if ci < len(row) else "", size=10)
            for rn in par.runs:
                if ri == 0:
                    rn.bold = True


lines = SRC.read_text(encoding="utf-8").split("\n")
i, first_heading, in_title = 0, True, False
while i < len(lines):
    raw = lines[i]
    s = raw.strip()
    if not s:
        i += 1
        continue
    if s.startswith(">") or s == "---":   # blockquote notu / yatay çizgi atlanır
        i += 1
        continue
    # Şekil
    mi = IMG.match(s)
    if mi:
        add_picture(mi.group(2).strip(), mi.group(1).strip())
        i += 1
        continue
    # Başlık
    mh = re.match(r"^(#{1,6})\s+(.*)$", s)
    if mh:
        lvl, txt = len(mh.group(1)), mh.group(2).strip()
        if first_heading:   # KAPAK BAŞLIĞI (Heading değil; ortalı büyük)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_run(p.add_run(txt), size=15, bold=True)
            first_heading, in_title = False, True
        else:
            in_title = False
            p = doc.add_paragraph(style=f"Heading {min(lvl, 3)}")
            p.add_run(txt)
            if lvl == 1 or txt in FRONT_BACK:
                p.paragraph_format.page_break_before = True
            if txt == "İÇİNDEKİLER":
                add_toc()
        i += 1
        continue
    # Tablo
    if s.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
            i += 1
        add_table(rows)
        continue
    # Liste
    lm = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", raw)
    if lm:
        marker, rest = lm.group(2), lm.group(3)
        buf = [rest]; i += 1
        while i < len(lines):
            nxt = lines[i]
            if nxt.strip() == "" or re.match(r"^(\s*)([-*]|\d+\.)\s+", nxt) or nxt.startswith("#") \
               or nxt.strip().startswith("|") or not nxt.startswith(" "):
                break
            buf.append(nxt.strip()); i += 1
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        style_run(p.add_run("• " if marker in ("-", "*") else marker + " "))
        add_inline(p, " ".join(buf))
        continue
    # Paragraf
    buf = [s]; i += 1
    while i < len(lines):
        nxt = lines[i]
        if nxt.strip() == "" or nxt.startswith("#") or nxt.strip().startswith("|") \
           or nxt.strip() in ("---",) or nxt.strip().startswith(">") \
           or re.match(r"^(\s*)([-*]|\d+\.)\s+", nxt) or IMG.match(nxt.strip()):
            break
        buf.append(nxt.strip()); i += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if in_title:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, " ".join(buf))

doc.save(str(OUT))
print("Kaydedildi ->", OUT)
print("Paragraf:", len(doc.paragraphs), "| Tablo:", len(doc.tables),
      "| Şekil:", len(doc.inline_shapes))
